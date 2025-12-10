"""
Semantic PAPL Parser
Extracts structured data from PAPL .docx documents

Outputs:
- pricing_data.json: Support items, prices, categories
- business_rules.yaml: Claiming rules, conditions, logic
- guidance.md: Policy text, explanations, requirements

ENHANCED: Now includes numerical price detection, comparison, and
table-structure anomaly detection for safer comparisons.
"""

from docx import Document
import json
import yaml
import os
from typing import Dict, List, Any, Tuple, Optional
import re
from datetime import datetime

# Import PDF page extractor if available
try:
    from pdf_page_extractor import PDFPageExtractor
    PDF_EXTRACTION_AVAILABLE = True
except ImportError:
    PDF_EXTRACTION_AVAILABLE = False
    PDFPageExtractor = None


class PAPLParser:
    """
    Parse PAPL documents into structured components

    Philosophy: PAPL contains three types of content:
    1. Pricing Data - Tables with support items and prices
    2. Business Rules - Conditions, thresholds, requirements
    3. Guidance - Explanatory text, policy context
    """

    def __init__(self):
        # print("🔥 USING NEW ROBUST PARSER 🔥")

        self.pricing_keywords = [
            "price",
            "rate",
            "cost",
            "fee",
            "charge",
            "amount",
            "support item",
            "support category",
            "registration group",
        ]

        self.rule_keywords = [
            "must",
            "should",
            "shall",
            "may not",
            "required",
            "mandatory",
            "condition",
            "threshold",
            "limit",
            "maximum",
            "minimum",
            "claiming",
            "quote",
            "evidence",
            "approval",
        ]

        self.guidance_keywords = [
            "guidance",
            "note",
            "example",
            "information",
            "consider",
            "background",
            "context",
            "purpose",
            "overview",
        ]

    # =====================================================================
    # PUBLIC PARSE ENTRY
    # =====================================================================
    def parse(self, docx_path_or_file, extract_page_numbers: bool = True) -> Dict[str, Any]:
        """
        Parse PAPL document into structured components

        Args:
            docx_path_or_file: Path to .docx or file-like object

        Returns:
            Dict with 'pricing_data', 'business_rules', 'guidance', 'metadata',
            plus 'raw_tables', 'raw_paragraphs', and 'anomalous_tables'.
        """
        doc = Document(docx_path_or_file)

        # Extract all content with metadata
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)  # ROBUST table extraction

        # Extract page numbers if requested and available
        if extract_page_numbers:
            if isinstance(docx_path_or_file, str) and os.path.exists(docx_path_or_file):
                try:
                    self._add_pdf_page_numbers(docx_path_or_file, tables)
                except Exception as e:
                    print(f"⚠️ PDF page extraction failed: {e}")
            else:
                print("⚠️ Cannot extract page numbers - need file path, not file object")

        # Detect anomalous tables (structurally unsafe for automated comparison)
        anomalous_tables: List[int] = []
        for t in tables:
            if self._detect_anomalous_structure(t):
                anomalous_tables.append(t["index"])
                t["is_anomalous"] = True
            else:
                t["is_anomalous"] = False

        # Classify and structure content
        pricing_data = self._extract_pricing_data(tables, paragraphs)
        business_rules = self._extract_business_rules(paragraphs, tables)
        guidance = self._extract_guidance(paragraphs)

        # Generate metadata
        metadata = self._generate_metadata(doc, paragraphs, tables)

        return {
            "pricing_data": pricing_data,
            "business_rules": business_rules,
            "guidance": guidance,
            "metadata": metadata,
            "raw_tables": tables,
            "raw_paragraphs": paragraphs,
            "anomalous_tables": anomalous_tables,
        }

    # =====================================================================
    # PARAGRAPHS + HEADINGS
    # =====================================================================
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract paragraphs with metadata"""
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            style_name = para.style.name if para.style else "Normal"
            paragraphs.append(
                {
                    "index": i,
                    "text": para.text.strip(),
                    "style": style_name,
                    "is_heading": "Heading" in style_name,
                    "level": self._get_heading_level(style_name),
                }
            )
        return paragraphs

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if not style_name or "Heading" not in style_name:
            return 0

        match = re.search(r"Heading\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
        return 0

    # =====================================================================
    # PAGE NUMBERS
    # =====================================================================
    def _add_pdf_page_numbers(self, docx_path: str, tables: List[Dict[str, Any]]) -> None:
        """Add exact page numbers to tables using PDF conversion"""
        print("📄 Extracting exact page numbers via PDF conversion...")

        try:
            extractor = PDFPageExtractor()

            if not extractor.is_available():
                print("⚠️ PDF extraction not available (missing LibreOffice or pdfplumber)")
                return

            table_pages = extractor.get_table_page_numbers(docx_path, len(tables))

            if table_pages:
                for table_idx, page_num in table_pages.items():
                    if table_idx < len(tables):
                        tables[table_idx]["page"] = page_num
                print(f"✓ Extracted page numbers for {len(table_pages)} tables")
            else:
                print("⚠️ Failed to extract page numbers from PDF")

        except Exception as e:
            print(f"⚠️ Error extracting page numbers: {e}")

    # =====================================================================
    # ROBUST TABLE EXTRACTION
    # =====================================================================
    def _normalize_cell_text(self, text: str) -> str:
        """Normalise cell text: strip, collapse whitespace, remove NBSP."""
        if text is None:
            return ""
        t = text.replace("\u00a0", " ")
        t = re.sub(r"\s+", " ", t)
        return t.strip()

    def _row_looks_like_data(self, row: List[str]) -> bool:
        """
        Heuristic: does this row look like a data row (contains item numbers / prices)?
        """
        for cell in row:
            if not cell:
                continue
            # Item number or obvious numeric price-ish values
            if self._extract_item_number(cell):
                return True
            if self._extract_price(cell):
                return True
            if re.search(r"\d{2}_\d{3}", cell):
                return True
            if re.search(r"\d+\.\d+", cell):
                return True
        return False

    def _infer_header_row_count(self, rows: List[List[str]]) -> int:
        """
        Infer how many top rows form the header block.
        Strategy:
          - Scan first few rows.
          - First row that "looks like data" marks the end of the header block.
        """
        if not rows:
            return 0

        max_header_rows = min(4, len(rows) - 1)  # leave at least one data row

        for i in range(max_header_rows):
            if self._row_looks_like_data(rows[i]):
                # Everything before this is header; if i==0, still treat as 1 header row
                return max(1, i)

        # If we never found a data-like row in first few, assume a single header row
        return 1

    def _collapse_header_rows(self, header_rows: List[List[str]]) -> List[str]:
        """
        Collapse multiple header rows into a single header row.
        - Concatenates non-empty texts in each column from top to bottom.
        - Fills empty columns by borrowing from neighbours where possible.
        """
        if not header_rows:
            return []

        num_cols = max(len(r) for r in header_rows)
        # Ensure all header rows have same length
        norm_rows = [r + [""] * (num_cols - len(r)) for r in header_rows]

        collapsed = []
        for col_idx in range(num_cols):
            parts: List[str] = []
            for r in norm_rows:
                t = r[col_idx].strip()
                if t and t not in parts:
                    parts.append(t)

            # Fallback: if empty, try neighbour columns
            if not parts:
                # look left then right
                for offset in (-1, 1, -2, 2):
                    neighbour = col_idx + offset
                    if 0 <= neighbour < num_cols:
                        for r in norm_rows:
                            t = r[neighbour].strip()
                            if t and t not in parts:
                                parts.append(t)
                        if parts:
                            break

            header = " ".join(parts).strip()
            if not header:
                header = f"col_{col_idx+1}"  # last resort
            collapsed.append(header)

        return collapsed

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract tables with robust grid + header handling:
          - Rectangular grid (rows padded to same length)
          - Multi-row headers collapsed into one header row
          - Headers stored in table_dict['headers']
        """
        tables: List[Dict[str, Any]] = []

        for i, table in enumerate(doc.tables):
            # Raw row extraction
            raw_rows: List[List[str]] = []
            for row in table.rows:
                cells = [self._normalize_cell_text(cell.text) for cell in row.cells]
                raw_rows.append(cells)

            if not raw_rows:
                tables.append(
                    {
                        "index": i,
                        "row_count": 0,
                        "col_count": 0,
                        "data": [],
                        "headers": [],
                        "is_pricing_table": False,
                        "pricing_confidence": 0,
                        "table_type": "other",
                    }
                )
                continue

            # Normalise all rows to same length
            max_cols = max(len(r) for r in raw_rows)
            norm_rows = [r + [""] * (max_cols - len(r)) for r in raw_rows]

            # Work out how many header rows
            header_row_count = self._infer_header_row_count(norm_rows)
            header_row_count = max(1, min(header_row_count, len(norm_rows) - 1))

            header_rows = norm_rows[:header_row_count]
            data_rows = norm_rows[header_row_count:]

            collapsed_headers = self._collapse_header_rows(header_rows)

            final_rows = [collapsed_headers] + data_rows

            table_dict: Dict[str, Any] = {
                "index": i,
                "row_count": len(final_rows),
                "col_count": max_cols,
                "data": final_rows,
                "headers": collapsed_headers,
            }

            # Pricing classification (uses new headers)
            is_pricing, confidence = self._is_pricing_table(table_dict)
            table_dict["is_pricing_table"] = is_pricing
            table_dict["pricing_confidence"] = confidence
            table_dict["table_type"] = self._classify_table_type(table_dict, is_pricing, confidence)

            tables.append(table_dict)

        return tables

    def _classify_table_type(self, table_dict: Dict[str, Any], is_pricing: bool, confidence: int) -> str:
        """Classify what type of table this is"""
        if is_pricing and confidence >= 60:
            return "pricing"

        headers = [str(h).lower() for h in table_dict.get("headers", [])]
        header_text = " ".join(headers)
        row_count = table_dict.get("row_count", 0) - 1  # exclude header

        # Example tables: few rows, very long narrative cells
        if row_count <= 5:
            data = table_dict.get("data", [])
            for row in data[1:]:
                for cell in row:
                    if len(str(cell)) > 500:
                        return "example"

        # Reference tables
        if any(keyword in header_text for keyword in ["mmm", "postcode", "location", "state", "zone"]):
            return "reference"

        # Metadata tables
        if any(keyword in header_text for keyword in ["version", "date published", "amendment", "page(s)"]):
            return "metadata"

        return "other"

    # =====================================================================
    # TABLE ANOMALY DETECTION
    # =====================================================================
    def _detect_anomalous_structure(self, table: Dict[str, Any]) -> bool:
        """
        Detect structural anomalies that make price comparison unsafe.

        A table is flagged as anomalous if any of the following hold:
        - Column headers contain state abbreviations (NSW, VIC, QLD, etc.)
        - More than 3 columns look like price/region columns
        - Column count is unusually large for a pricing table (> 6)

        This is deliberately conservative: better to suppress comparison
        than to report misleading price deltas off mismatched structures.
        """
        headers = [str(h).lower() for h in table.get("headers", [])]

        # Rule 1: state abbreviations in headers (e.g. "nsw vic qld act")
        state_tokens = ["nsw", "vic", "qld", "act", "wa", "sa", "tas", "nt"]
        for h in headers:
            for state in state_tokens:
                if state in h:
                    return True

        # Rule 2: more than 3 price-like/region-like columns
        price_like_keywords = [
            "national",
            "remote",
            "very remote",
            "price",
            "rate",
            "cost",
            "limit",
            "amount",
            "fee",
            "charge",
        ]
        price_like_count = 0
        for h in headers:
            if any(k in h for k in price_like_keywords):
                price_like_count += 1
        if price_like_count > 3:
            return True

        # Rule 3: suspiciously wide table (more than 6 columns)
        if table.get("col_count", 0) > 6:
            return True

        return False

    # =====================================================================
    # PRICING DATA EXTRACTION
    # =====================================================================
    def _extract_pricing_data(self, tables: List[Dict], paragraphs: List[Dict]) -> Dict[str, Any]:
        """Extract pricing data from tables"""
        pricing_data = {
            "support_items": [],
            "categories": set(),
            "total_items": 0,
        }

        for table in tables:
            if not table["data"] or len(table["data"]) < 2:
                continue

            headers = [str(h).lower() for h in table["data"][0]]
            has_pricing = any(keyword in " ".join(headers) for keyword in self.pricing_keywords)

            if has_pricing:
                for row_idx, row in enumerate(table["data"][1:], start=1):
                    if len(row) < 2:
                        continue

                    item = {
                        "table_index": table["index"],
                        "row_index": row_idx,
                        "data": row,
                        "price": self._extract_price(row),
                        "item_number": self._extract_item_number(" ".join(str(c) for c in row)),
                    }
                    pricing_data["support_items"].append(item)

        pricing_data["total_items"] = len(pricing_data["support_items"])
        pricing_data["categories"] = list(pricing_data["categories"])
        return pricing_data

    # =====================================================================
    # BUSINESS RULES / GUIDANCE
    # =====================================================================
    def _extract_business_rules(self, paragraphs: List[Dict], tables: List[Dict]) -> Dict[str, Any]:
        """Extract business rules from paragraphs"""
        rules = {
            "claiming_rules": [],
            "conditions": [],
            "thresholds": [],
            "total_rules": 0,
        }

        for para in paragraphs:
            text_lower = para["text"].lower()
            has_rule_keywords = any(keyword in text_lower for keyword in self.rule_keywords)

            if has_rule_keywords:
                rule = {
                    "paragraph_index": para["index"],
                    "text": para["text"],
                    "type": self._classify_rule(para["text"]),
                }

                if "claiming" in text_lower or "claim" in text_lower:
                    rules["claiming_rules"].append(rule)
                elif any(word in text_lower for word in ["condition", "if", "when"]):
                    rules["conditions"].append(rule)
                elif any(word in text_lower for word in ["threshold", "limit", "maximum", "minimum"]):
                    rules["thresholds"].append(rule)

        rules["total_rules"] = (
            len(rules["claiming_rules"]) + len(rules["conditions"]) + len(rules["thresholds"])
        )
        return rules

    def _extract_guidance(self, paragraphs: List[Dict]) -> Dict[str, Any]:
        """Extract guidance text"""
        guidance = {"sections": [], "markdown": "", "total_paragraphs": 0}
        current_section = None

        for para in paragraphs:
            if para["is_heading"]:
                if current_section:
                    guidance["sections"].append(current_section)
                current_section = {"title": para["text"], "level": para["level"], "paragraphs": []}
            else:
                text_lower = para["text"].lower()
                has_guidance = any(keyword in text_lower for keyword in self.guidance_keywords)
                if has_guidance or (current_section and len(para["text"]) > 50):
                    if current_section:
                        current_section["paragraphs"].append(para["text"])

        if current_section:
            guidance["sections"].append(current_section)

        guidance["markdown"] = self._generate_markdown(guidance["sections"])
        guidance["total_paragraphs"] = sum(len(s["paragraphs"]) for s in guidance["sections"])
        return guidance

    def _classify_rule(self, text: str) -> str:
        """Classify type of rule"""
        text_lower = text.lower()
        if "must" in text_lower or "shall" in text_lower or "required" in text_lower:
            return "mandatory"
        elif "should" in text_lower or "recommended" in text_lower:
            return "recommended"
        elif "may" in text_lower or "can" in text_lower:
            return "optional"
        else:
            return "informational"

    def _generate_metadata(self, doc: Document, paragraphs: List[Dict], tables: List[Dict]) -> Dict[str, Any]:
        """Generate document metadata"""
        return {
            "parsed_at": datetime.now().isoformat(),
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables),
            "total_headings": sum(1 for p in paragraphs if p["is_heading"]),
            "document_length": sum(len(p["text"]) for p in paragraphs),
        }

    def _generate_markdown(self, sections: List[Dict]) -> str:
        """Generate markdown from guidance sections"""
        md_lines: List[str] = []
        for section in sections:
            heading_prefix = "#" * section["level"] if section["level"] > 0 else "##"
            md_lines.append(f"{heading_prefix} {section['title']}\n")
            for para in section["paragraphs"]:
                md_lines.append(f"{para}\n")
            md_lines.append("")
        return "\n".join(md_lines)

    # =====================================================================
    # EXPORTS
    # =====================================================================
    def export_to_json(self, parsed_data: Dict[str, Any]) -> str:
        export_data = {"pricing_data": parsed_data["pricing_data"], "metadata": parsed_data["metadata"]}
        return json.dumps(export_data, indent=2)

    def export_to_yaml(self, parsed_data: Dict[str, Any]) -> str:
        export_data = {
            "business_rules": parsed_data["business_rules"],
            "metadata": {
                "parsed_at": parsed_data["metadata"]["parsed_at"],
                "total_rules": parsed_data["business_rules"]["total_rules"],
            },
        }
        return yaml.dump(export_data, default_flow_style=False, sort_keys=False)

    def export_to_markdown(self, parsed_data: Dict[str, Any]) -> str:
        return parsed_data["guidance"]["markdown"]

    # =====================================================================
    # PRICE / ITEM EXTRACTION HELPERS
    # =====================================================================
    def _extract_price(self, item_or_row) -> Optional[float]:
        """Robust price extractor supporting:
        - strings
        - lists of cells
        - dicts
        - any nested structure
        """

        # Normalise to a string
        if isinstance(item_or_row, list):
            text = " ".join(str(x) for x in item_or_row)
        elif isinstance(item_or_row, dict):
            text = " ".join(str(v) for v in item_or_row.values())
        else:
            text = str(item_or_row)

        if not text:
            return None

        # Remove currency symbols
        cleaned = text.replace("$", "").replace("€", "").replace("£", "").strip()

        # Remove thousands separators
        cleaned = cleaned.replace(",", "")

        # Extract any integer or decimal number
        match = re.search(r"-?\d+(?:\.\d+)?", cleaned)
        if not match:
            return None

        try:
            return round(float(match.group(0)), 2)
        except ValueError:
            return None

    def _extract_all_prices(self, text) -> List[float]:
        if isinstance(text, list):
            text = " ".join(str(x) for x in text)
        elif isinstance(text, dict):
            text = " ".join(str(v) for v in text.values())
        else:
            text = str(text)

        cleaned = text.replace("$", "").replace("€", "").replace("£", "")
        cleaned = cleaned.replace(",", "")

        matches = re.findall(r"-?\d+(?:\.\d+)?", cleaned)

        prices = []
        for m in matches:
            try:
                prices.append(round(float(m), 2))
            except ValueError:
                continue

        return prices

    def _extract_item_number(self, text: str) -> Optional[str]:
        """Extract NDIS support item number from text"""
        if not text:
            return None
        patterns = [
            r"(\d{2}_\d{3}_\d{4}_\d+_\d+)",
            r"(\d{2}_\d{3}_\d{4})",
            r"(\d{2}_\d{3})",
            r"(\d+\.\d+\.\d+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, str(text))
            if match:
                return match.group(1)
        return None

    # =====================================================================
    # PRICING TABLE CLASSIFICATION
    # =====================================================================
    def _is_pricing_table(
        self, table: Dict[str, Any], context_paragraphs: List[Dict[str, Any]] = None
    ) -> Tuple[bool, int]:
        confidence_score = 0
        data = table.get("data", [])
        if not data or len(data) < 2:
            return False, 0

        headers = [str(h).lower() for h in data[0]]
        header_text = " ".join(headers)
        row_count = len(data) - 1

        # Table size
        if row_count >= 50:
            confidence_score += 40
        elif row_count >= 20:
            confidence_score += 30
        elif row_count >= 10:
            confidence_score += 20

        # Price columns
        price_column_keywords = [
            "national",
            "remote",
            "very remote",
            "price",
            "rate",
            "cost",
            "limit",
            "amount",
        ]
        price_columns = sum(1 for h in headers if any(keyword in h for keyword in price_column_keywords))

        if price_columns >= 3:
            confidence_score += 40
        elif price_columns >= 2:
            confidence_score += 30
        elif price_columns >= 1:
            confidence_score += 20

        # Item number patterns
        item_number_count = 0
        if len(data) > 1:
            for row in data[1 : min(6, len(data))]:
                if row:
                    cell = str(row[0])
                    if self._extract_item_number(cell):
                        item_number_count += 1
            if item_number_count >= 3:
                confidence_score += 30

        if any(keyword in header_text for keyword in ["item number", "support item", "item name"]):
            confidence_score += 10

        if "unit" in header_text:
            confidence_score += 10

        if context_paragraphs:
            context_text = " ".join(p.get("text", "").lower() for p in context_paragraphs[-3:])
            if any(word in context_text for word in ["pricing", "price", "rate", "cost"]):
                confidence_score += 10

        if price_columns > 0 and row_count > 0:
            price_cell_count = 0
            for row in data[1 : min(6, len(data))]:
                for cell in row:
                    if self._extract_price(str(cell)):
                        price_cell_count += 1
            if price_cell_count >= 5:
                confidence_score += 10

        # Penalise very long narrative cells
        for row in data[1 : min(6, len(data))]:
            for cell in row:
                if len(str(cell)) > 500:
                    confidence_score -= 30
                    break

        confidence_score = max(0, confidence_score)
        return confidence_score >= 60, confidence_score

    # =====================================================================
    # PER-TABLE PRICE EXTRACTION + COMPARISON
    # =====================================================================
    def _extract_prices_from_table(self, table: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract all prices from a table with location info"""
        prices: List[Dict[str, Any]] = []
        if not table.get("data"):
            return prices

        table_index = table.get("index", 0)

        for row_idx, row in enumerate(table["data"][1:], start=1):
            item_number = None
            item_description = None

            for cell in row:
                cell_text = str(cell)
                if not item_number:
                    item_number = self._extract_item_number(cell_text)
                if not item_description and len(cell_text) > 10 and not self._extract_price(cell_text):
                    item_description = cell_text.strip()

            for col_idx, cell in enumerate(row):
                if col_idx == 0:
                    continue
                cell_text = str(cell)
                price = self._extract_price(cell_text)
                if price:
                    prices.append(
                        {
                            "table_index": table_index,
                            "item_number": item_number,
                            "item_description": item_description,
                            "price": price,
                            "row_index": row_idx,
                            "col_index": col_idx,
                            "raw_text": cell_text.strip(),
                            "page": table.get("page", 0),
                        }
                    )

        return prices

    def _compare_table_prices(self, old_table: Dict[str, Any], new_table: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Compare prices between two tables by item number.

        If either table is flagged as structurally anomalous, we suppress
        automated comparison and return an empty list. This avoids
        misleading % changes when old/new structures don't align
        (e.g. state-grouped columns vs national/remote/very remote).
        """
        # Skip comparison if table structure is unsafe
        if self._detect_anomalous_structure(old_table) or self._detect_anomalous_structure(new_table):
            return []

        changes: List[Dict[str, Any]] = []
        old_prices = self._extract_prices_from_table(old_table)
        new_prices = self._extract_prices_from_table(new_table)

        old_by_item_col: Dict[Tuple[str, int], Dict[str, Any]] = {}
        for p in old_prices:
            if p["item_number"]:
                key = (p["item_number"], p["col_index"])
                old_by_item_col[key] = p

        new_by_item_col: Dict[Tuple[str, int], Dict[str, Any]] = {}
        for p in new_prices:
            if p["item_number"]:
                key = (p["item_number"], p["col_index"])
                new_by_item_col[key] = p

        for key, old_p in old_by_item_col.items():
            if key in new_by_item_col:
                new_p = new_by_item_col[key]
                item_number, col_index = key

                if old_p["price"] != new_p["price"]:
                    difference = new_p["price"] - old_p["price"]
                    percent_change = (difference / old_p["price"]) * 100 if old_p["price"] > 0 else 0

                    changes.append(
                        {
                            "item_number": item_number,
                            "item_description": new_p["item_description"] or old_p["item_description"],
                            "old_price": old_p["price"],
                            "new_price": new_p["price"],
                            "difference": difference,
                            "percent_change": percent_change,
                            "old_location": {
                                "table": old_p["table_index"],
                                "row": old_p["row_index"],
                                "col": old_p["col_index"],
                                "page": old_p.get("page", 0),
                            },
                            "new_location": {
                                "table": new_p["table_index"],
                                "row": new_p["row_index"],
                                "col": new_p["col_index"],
                                "page": new_p.get("page", 0),
                            },
                            "old_raw_text": old_p["raw_text"],
                            "new_raw_text": new_p["raw_text"],
                        }
                    )

        return changes


if __name__ == "__main__":
    parser = PAPLParser()
    print("PAPL Parser ready with robust table extraction, anomaly detection, and price comparison!")
